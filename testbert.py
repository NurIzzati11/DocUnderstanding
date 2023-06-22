##all libraries used
import streamlit as st
import pdfplumber
import docx
import numpy 
import pandas as pd
import nltk
import re
import gensim
import gensim.downloader
##from gensim.parsing.preprocessing import preprocess_string
#pd.set_option('display.max_colwidth', None)
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.neighbors import NearestNeighbors
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from gensim.parsing.preprocessing import remove_stopwords
import nltk
#nltk.download('punkt')
from sklearn.metrics.pairwise import cosine_similarity
from gensim.models import Word2Vec
import gensim.downloader as api
from haystack.document_stores import InMemoryDocumentStore
from haystack.nodes import BM25Retriever
from haystack.nodes import FARMReader
from haystack.pipelines import ExtractiveQAPipeline
from haystack.utils import print_answers
from haystack import Document    
    
##file collecting
documents = []
bboxes=[]

def readDoc(filename):
    doc = docx.Document(filename)
    fullText = ''
    for para in doc.paragraphs:
        fullText.append(para.text)
    documents.append(Document(content=para.text))
    return ' '.join(fullText)


def readPDF(filename):
    pdf = pdfplumber.open(filename)
    all_texts = ""

    for i in range(0 ,len(pdf.pages)):
        page = pdf.pages[i]
        bboxes = [table.bbox for table in page.find_tables()]
        text = page.filter(not_within_bboxes).extract_text(x_tolerance=1, y_tolerance=0)
        text = text.replace("\n"," ")
        documents.append(Document(content=text))
        all_texts += text

    return all_texts
  
def not_within_bboxes(obj):
    def obj_in_bbox(_bbox):
        v_mid = (obj["top"] + obj["bottom"]) / 2
        h_mid = (obj["x0"] + obj["x1"]) / 2
        x0, top, x1, bottom = _bbox
        return (h_mid >= x0) and (h_mid < x1) and (v_mid >= top) and (v_mid < bottom)
    return not any(obj_in_bbox(__bbox) for __bbox in bboxes)
    
#STREAMLIT
st.markdown("<h1 style='text-align: center; color: white;'>DOCUMENT UNDERSTANDING ANALYZER</h1>", unsafe_allow_html=True)
st.markdown("<h6 style='text-align: center; color: white;'>Upload a document and start your Q&A!</h6>", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Choose a file (PDF or DOCX only)")
if uploaded_file is not None:
    filename=uploaded_file.name
    if filename.endswith('.docx'):
        text = readDoc(uploaded_file)
    elif filename.endswith('.pdf'):
        text = readPDF(uploaded_file)

#User question
user_input = st.text_area(label='Type your question here....')
st.write("Choose your models")
c1,c2,c3,c4,c5=st.columns(5)
with c1:
    bert=st.button("BERT")
    
if bert==1:
    st.write("BERT Model answers")
    with st.spinner("Searching. Please hold..."):
    ##BERT
        document_store = InMemoryDocumentStore(use_bm25=True)
        document_store.delete_documents()
        document_store.write_documents(documents)
        retriever = BM25Retriever(document_store=document_store)
        reader = FARMReader(model_name_or_path="deepset/roberta-base-squad2", use_gpu=True)
        pipe = ExtractiveQAPipeline(reader, retriever)
        res_br = pipe.run(query=user_input, params={"Retriever": {"top_k": 5},"Reader": {"top_k": 5}})
        st.write("Here are the top 3 answers")
        df = pd.DataFrame(columns=['Number', 'Answer', 'Score'])
        for i in range(0,3):
            ans=res_br['answers'][i].answer
            score=res_br['answers'][i].score
            df = df.concat({'Rank': i+1, 'Answer': ans, 'Score': score}, ignore_index=True)
            #st.write("Answer " + str(i+1) + ": " + ans + " (Score: " + str(score) + ")")
        st.dataframe(df,hide_index=True)    
      
