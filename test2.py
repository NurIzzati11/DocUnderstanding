##all libraries used
import streamlit as st
import pdfplumber
import docx
import numpy 
import pandas as pd
import nltk
import gensim
import gensim.downloader
from gensim.parsing.preprocessing import preprocess_string
pd.set_option('display.max_colwidth', None)
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.neighbors import NearestNeighbors
from sklearn.metrics.pairwise import cosine_similarity
from nltk.corpus import stopwords
from gensim.parsing.preprocessing import remove_stopwords
import nltk
from sklearn.metrics.pairwise import cosine_similarity
from gensim.models import Word2Vec
import gensim.downloader as api
from haystack.document_stores import InMemoryDocumentStore
from haystack.nodes import BM25Retriever
from haystack.nodes import FARMReader
from haystack.pipelines import ExtractiveQAPipeline
from haystack.utils import print_answers
from haystack import Document    
    
    
##file collecting
documents = []

def readDoc(filename):
    doc = docx.Document(filename)
    fullText = ''
    for para in doc.paragraphs:
        fullText.append(para.text)
    documents.append(Document(content=para.text))
    return ' '.join(fullText)

def not_within_bboxes(obj):
    def obj_in_bbox(_bbox):
        v_mid = (obj["top"] + obj["bottom"]) / 2
        h_mid = (obj["x0"] + obj["x1"]) / 2
        x0, top, x1, bottom = _bbox
        return (h_mid >= x0) and (h_mid < x1) and (v_mid >= top) and (v_mid < bottom)
    return not any(obj_in_bbox(__bbox) for __bbox in bboxes)


def readPDF(filename):
    pdf = pdfplumber.open(filename)
    all_texts = ""

    for i in range(0 ,len(pdf.pages)):
        page = pdf.pages[i]
        bboxes = [table.bbox for table in page.find_tables()]
        text = page.filter(not_within_bboxes).extract_text(x_tolerance=1, y_tolerance=0)
        text = text.replace("\n"," ")
        documents.append(Document(content=text))
        all_texts += text

    return all_texts

#Clean the texts and other functions
def clean_sentence(sentence, stopwords=False):
  sentence = sentence.lower().strip()
  sentence = re.sub(r'[^a-z0-9\s]', '', sentence)
  if stopwords:
    sentence = remove_stopwords(sentence)
  return sentence

def get_cleaned_sentences(tokens, stopwords=False):
  cleaned_sentences = []
  for row in tokens:
    cleaned = clean_sentence(row, stopwords)
    cleaned_sentences.append(cleaned)
  return cleaned_sentences

def retrieveAndPrintFAQAnswer(question_embedding, sentence_embeddings, sentences):
  max_sim = -1
  index_sim = -1
  for index, embedding in enumerate(sentence_embeddings):
    sim = cosine_similarity(embedding, question_embedding)[0][0]
    # print(index, sim, sentences[index])
    if sim > max_sim:
      max_sim = sim
      index_sim = index
  return index_sim

def getWordVec(word, model):
        samp = model['pc']
        vec = [0]*len(samp)
        try:
            vec = model[word]
        except:
            vec = [0]*len(samp)
        return (vec)


def getPhraseEmbedding(phrase, embeddingmodel):
  samp = getWordVec('computer', embeddingmodel)
  vec = numpy.array([0]*len(samp))
  den = 0;
  for word in phrase.split():
    den = den+1
    vec = vec+numpy.array(getWordVec(word, embeddingmodel))
  return vec.reshape(1, -1)

#model import
v2w_model = None
try:
    v2w_model = gensim.models.Keyedvectors.load('./w2vecmodel.mod')
except:
    v2w_model = api.load('word2vec-google-news-300')
    v2w_model.save("./w2vecmodel.mod")
        
#STREAMLIT
st.title("Document Understanding Analyzer")

##Intro Part:
st.write("""Introduction and User Manual:
         \n......
         \nPlease choose your file to start!
         """)

#Upload file
file = st.file_uploader("Choose a file (PDF or DOCX only)")
if file.lower().endswith('.docx'):
    text = readDoc(file)
elif file.lower().endswith('.pdf'):
    text = readPDF(file)
else: print('Only pdf and docx are accepted')

#User question
user_input = st.text_area(label='Type your question here....')
user_click=st.button('Submit')
tabs= st.tabs(["Word2Vec", "BERT"])
if user_click==1:
    
    ##WORD2VEC
    tab_word2v = tabs[0]
    tab_word2v.subheader("WORD2VEC")
    tab_word2v.write("Simple ML Model")
    
    tokens = nltk.sent_tokenize(text)
    cleaned_sentences = get_cleaned_sentences(tokens, stopwords=True)
    cleaned_sentences_with_stopwords = get_cleaned_sentences(tokens, stopwords=False)
    sentences = cleaned_sentences_with_stopwords

    sent_embeddings = []
    for sent in sentences:
        sent_embeddings.append(getPhraseEmbedding(sent, v2w_model))

    question = clean_sentence(user_input, stopwords=False)
    question_embedding = getPhraseEmbedding(question, v2w_model)
    index = retrieveAndPrintFAQAnswer(question_embedding, sent_embeddings, cleaned_sentences_with_stopwords)
    res_w2v = cleaned_sentences_with_stopwords[index]
    tab_word2v.write(res_w2v)
    
    ##BERT
    tab_bert = tabs[1]
    tab_bert.subheader("BERT")
    tab_bert.write("Advanced ML Model")
    
    document_store = InMemoryDocumentStore(use_bm25=True)
    document_store.delete_documents()
    document_store.write_documents(documents)
    retriever = BM25Retriever(document_store=document_store)
    reader = FARMReader(model_name_or_path="deepset/roberta-base-squad2", use_gpu=True)
    pipe = ExtractiveQAPipeline(reader, retriever)
    res_br = pipe.run(query=user_input, params={"Retriever": {"top_k": 5},"Reader": {"top_k": 5}})
    tab_bert.write(print_answers(res_br, details="minimum"))



